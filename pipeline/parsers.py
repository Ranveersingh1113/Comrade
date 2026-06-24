"""Document parsers + spotlighting for the pipeline.

Each parser returns plain/markdown text and is pure (bytes/str in, str out).
We extract only visible body content and do not pass through document metadata,
so hidden metadata can't leak into the model. All extracted text is untrusted
and must be passed through spotlight() before it reaches the compiler's model.
"""
import io
import re

import pymupdf
import pymupdf4llm
from docx import Document

# datamarking marker: spaces in untrusted text become this, and the compiler's
# system prompt declares it, so the model treats marked spans as data not
# instructions. (Microsoft Research: datamarking > delimiters, low task impact.)
SPACE_MARK = "^"


def parse_pdf(data: bytes) -> str:
    """PDF bytes -> Markdown (reading order, tables) via PyMuPDF4LLM."""
    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        return pymupdf4llm.to_markdown(doc).strip()
    finally:
        doc.close()


def parse_docx(data: bytes) -> str:
    """.docx bytes -> text (paragraphs + table cells), visible body only."""
    document = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


# matches a WhatsApp line prefix: "[31/12/2023, 23:59:59] " or "12/31/23, 11:59 PM - "
_WA_PREFIX = re.compile(
    r"^\[?\d{1,2}[/.]\d{1,2}[/.]\d{2,4},?\s+"
    r"\d{1,2}:\d{2}(?::\d{2})?\s*(?:[APap][Mm])?\]?\s*[-–]?\s*"
)


def parse_whatsapp(text: str) -> str:
    """WhatsApp .txt export -> normalised "Sender: message" lines.

    Lines without a sender (system notices like the encryption banner) are
    dropped; lines without a timestamp prefix are treated as continuations.
    """
    out: list[str] = []
    for raw in text.splitlines():
        m = _WA_PREFIX.match(raw)
        if m:
            rest = raw[m.end():]
            if ":" in rest:  # has "Sender: body"; otherwise a system notice
                sender, _, body = rest.partition(":")
                out.append(f"{sender.strip()}: {body.strip()}")
        elif out and raw.strip():
            out[-1] += " " + raw.strip()  # continuation of previous message
    return "\n".join(out)


def spotlight(text: str) -> str:
    """Datamark untrusted text: replace spaces with SPACE_MARK. Newlines kept so
    structure survives. The compiler declares this marking in its system prompt."""
    return text.replace(" ", SPACE_MARK)
